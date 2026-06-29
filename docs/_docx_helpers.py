"""
Helpers compartidos para construir los informes Word (EP3 y EFT) con python-docx.
Define estilos corporativos, encabezados, tablas con formato y utilidades de
inserción de figuras a partir de los datos reales de observabilidad.
"""

from __future__ import annotations

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHARTS = os.path.join(ROOT, "dashboard", "charts")

AZUL = RGBColor(0x1F, 0x4E, 0x79)
AZUL_FILL = "1F4E79"
GRIS_FILL = "D9E1F2"
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)


# ── Configuración base del documento ──────────────────────────────────────────

def nuevo_documento() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    # Estilos de encabezado en azul corporativo
    for nivel, size in (("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)):
        st = doc.styles[nivel]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = AZUL
        st.font.bold = True
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)

    # Márgenes 2 cm
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)
    return doc


def set_page_numbers(doc: Document, texto_izq: str):
    """Pie de página con texto a la izquierda y 'Página X' a la derecha."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = ""
    run = p.add_run(texto_izq + "\t\t")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    # Campo PAGE
    run2 = p.add_run("Página ")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    _add_field(p, "PAGE")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Tabs para alinear el número a la derecha
    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.8), WD_TAB_ALIGNMENT.RIGHT)


def _add_field(paragraph, field_code: str):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)


# ── Bloque de título ──────────────────────────────────────────────────────────

def bloque_titulo(doc, titulo, subtitulo, sigla, autores, fecha, ponderacion=None):
    # Barra superior
    p = doc.add_paragraph()
    r = p.add_run(sigla)
    r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = AZUL
    if ponderacion:
        r2 = p.add_run(f"   ·   {ponderacion}")
        r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    t = doc.add_paragraph()
    rt = t.add_run(titulo)
    rt.font.size = Pt(20); rt.font.bold = True; rt.font.color.rgb = AZUL
    t.paragraph_format.space_after = Pt(2)

    s = doc.add_paragraph()
    rs = s.add_run(subtitulo)
    rs.font.size = Pt(12); rs.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    s.paragraph_format.space_after = Pt(6)

    meta = doc.add_paragraph()
    rm = meta.add_run(f"Autores: {autores}\nFecha: {fecha}")
    rm.font.size = Pt(9.5); rm.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    _barra(doc)


def _barra(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), AZUL_FILL)
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(6)


# ── Párrafos y listas ─────────────────────────────────────────────────────────

def h1(doc, texto): doc.add_heading(texto, level=1)
def h2(doc, texto): doc.add_heading(texto, level=2)
def h3(doc, texto): doc.add_heading(texto, level=3)


def parrafo(doc, texto, size=10.5, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(size); r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rico(doc, segmentos, size=10.5, space_after=6):
    """Párrafo con segmentos (texto, bold) para resaltar términos."""
    p = doc.add_paragraph()
    for texto, bold in segmentos:
        r = p.add_run(texto); r.font.size = Pt(size); r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def vinneta(doc, items, size=10.5):
    for it in items:
        if isinstance(it, str):
            p = doc.add_paragraph(it, style="List Bullet")
            p.runs[0].font.size = Pt(size)
        else:  # lista/tupla de segmentos (texto, bold)
            p = doc.add_paragraph(style="List Bullet")
            for texto, bold in it:
                r = p.add_run(texto); r.font.size = Pt(size); r.bold = bold
        p.paragraph_format.space_after = Pt(2)


def numerada(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        if isinstance(it, list):
            for texto, bold in it:
                r = p.add_run(texto); r.font.size = Pt(size); r.bold = bold
        else:
            r = p.add_run(it); r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


# ── Tablas ────────────────────────────────────────────────────────────────────

def tabla(doc, encabezados, filas, anchos=None, size=9, fuente_resaltar_col0=False):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, enc in enumerate(encabezados):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(enc)
        run.font.bold = True; run.font.size = Pt(size); run.font.color.rgb = BLANCO
        _shade(hdr[i], AZUL_FILL)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, fila in enumerate(filas):
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(size)
            if fuente_resaltar_col0 and i == 0:
                run.font.bold = True
            celdas[i].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            if j % 2 == 1:
                _shade(celdas[i], GRIS_FILL)
    if anchos:
        for row in t.rows:
            for i, w in enumerate(anchos):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# ── Figuras ───────────────────────────────────────────────────────────────────

def figura(doc, nombre_png, caption, width=5.6):
    ruta = os.path.join(CHARTS, nombre_png)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(ruta, width=Inches(width))
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(caption)
    rc.font.size = Pt(8.5); rc.italic = True; rc.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    cap.paragraph_format.space_after = Pt(6)


def caja(doc, titulo, texto):
    """Recuadro destacado (callout) con sombreado suave."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    _shade(cell, "EAF1FB")
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(titulo); r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AZUL
    p2 = cell.add_paragraph()
    r2 = p2.add_run(texto); r2.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def salto_pagina(doc):
    from docx.enum.text import WD_BREAK
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
