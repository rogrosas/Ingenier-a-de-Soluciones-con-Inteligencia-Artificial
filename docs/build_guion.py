"""
Guion de la Presentación (Word) — diapositiva por diapositiva, para la defensa EFT.
Acompaña a EFT_Presentacion_Defensa.pptx. Incluye tiempos, qué mostrar, qué decir
(guía a personalizar) y transiciones, más un bloque de preguntas frecuentes.
"""

from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx.shared import Pt, Inches, RGBColor

from docs._docx_helpers import (
    ROOT, nuevo_documento, bloque_titulo, set_page_numbers,
    h1, h2, parrafo, rico, vinneta, tabla, caja,
)

ENTREGA = os.path.join(ROOT, "entrega"); os.makedirs(ENTREGA, exist_ok=True)
SALIDA = os.path.join(ENTREGA, "GUION_PRESENTACION_EFT.docx")
ROJO = RGBColor(0xC0, 0x39, 0x2B)
GRIS = RGBColor(0x59, 0x59, 0x59)
AZUL = RGBColor(0x1F, 0x4E, 0x79)


def nota_ia(doc, texto):
    t = doc.add_paragraph()
    r = t.add_run("⚠ " + texto)
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = ROJO
    t.paragraph_format.space_after = Pt(6)


def slide(doc, n, titulo, tiempo, mostrar, decir, transicion=None):
    h2(doc, f"Diapositiva {n} — {titulo}   ·   {tiempo}")
    p = doc.add_paragraph()
    r = p.add_run("Qué mostrar: "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUL
    r2 = p.add_run(mostrar); r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    pd = doc.add_paragraph()
    rd = pd.add_run("Qué decir: "); rd.bold = True; rd.font.size = Pt(10); rd.font.color.rgb = AZUL
    rdt = pd.add_run("“" + decir + "”"); rdt.font.size = Pt(10); rdt.italic = True
    pd.paragraph_format.space_after = Pt(3)
    if transicion:
        pt = doc.add_paragraph()
        rt = pt.add_run("Transición: "); rt.bold = True; rt.font.size = Pt(9); rt.font.color.rgb = GRIS
        rtt = pt.add_run(transicion); rtt.font.size = Pt(9); rtt.font.color.rgb = GRIS
        pt.paragraph_format.space_after = Pt(8)


def build():
    doc = nuevo_documento()
    set_page_numbers(doc, "Guion de Defensa · Banco Digital Chile · ISY0101")

    bloque_titulo(
        doc,
        titulo="Guion de la Presentación",
        subtitulo="Defensa EFT · Agente “Banco Digital Chile” · 10 minutos",
        sigla="ISY0101 · Ingeniería de Soluciones con IA",
        autores="Roger Rosas · Rodrigo Santis · Edgardo Gutiérrez",
        fecha="Junio 2025",
        ponderacion="Acompaña a EFT_Presentacion_Defensa.pptx",
    )

    parrafo(doc,
        "Este guion mapea 1:1 con las 13 diapositivas. La defensa simula una reunión "
        "ejecutiva: el/la docente representa a la dirección. Habla con seguridad, mira a la "
        "audiencia y apóyate en evidencia. Practica al menos una vez de principio a fin.")
    nota_ia(doc,
        "El texto «qué decir» es una GUÍA para que la adaptes a tu voz. Las justificaciones de "
        "diseño y las reflexiones deben ser propias (la pauta prohíbe leer respuestas generadas "
        "por IA en esas partes).")

    h1(doc, "Distribución del tiempo (≈10 min)")
    tabla(doc, ["#", "Diapositiva", "Tiempo", "Foco rúbrica"], [
        ["1", "Portada", "0:20", "Apertura"],
        ["2", "El caso", "0:45", "IE1, IE4"],
        ["3", "La solución", "0:40", "IE4, IE8"],
        ["4", "Arquitectura", "0:45", "IE3, IE8"],
        ["5", "Agente funcional", "1:15", "IE5, IE6, IE7"],
        ["6", "Evidencia — pipeline y pruebas", "0:45", "IE9"],
        ["7", "Evidencia — seguridad y RAG", "0:45", "IE11, IE2"],
        ["8", "Diseño LLM + RAG", "0:40", "IE1, IE2"],
        ["9", "Observabilidad", "1:05", "IE9"],
        ["10", "Trazabilidad", "0:55", "IE10, IE12"],
        ["11", "Seguridad y ética", "0:45", "IE11"],
        ["12", "Mejoras", "0:40", "IE12"],
        ["13", "Cierre", "0:15", "IE13–IE15"],
    ], anchos=[0.4, 3.1, 1.0, 1.9], size=9.5, fuente_resaltar_col0=True)

    h1(doc, "Guion diapositiva por diapositiva")

    slide(doc, 1, "Portada", "0:20",
        "Título del proyecto y tu nombre.",
        "Buenos días. Soy [tu nombre], del equipo que diseñó un agente de inteligencia "
        "artificial para Banco Digital Chile. En los próximos diez minutos les mostraré el caso, "
        "la solución, una demostración y los resultados de su monitoreo.",
        "Avanza al caso.")

    slide(doc, 2, "El caso organizacional", "0:45",
        "Viñetas del problema y el recuadro del desafío.",
        "El banco opera de forma 100% digital y recibe un alto volumen de consultas repetitivas "
        "y operaciones sensibles. La atención humana no escala y un chatbot de reglas no entiende "
        "lenguaje natural ni ejecuta acciones. El desafío era automatizar la atención de extremo "
        "a extremo, sin perder seguridad ni trazabilidad.",
        "“Frente a esto, propusimos lo siguiente.”")

    slide(doc, 3, "La solución propuesta", "0:45",
        "Las cinco tarjetas de capacidades.",
        "Construimos un agente conversacional que no solo responde: consulta saldos y productos, "
        "ejecuta transferencias y reclamos, razona simulaciones de crédito, recuerda al cliente "
        "entre sesiones y coordina especialistas. Es decir, resuelve la tarea completa.",
        "“¿Cómo se organiza por dentro? Esta es la arquitectura.”")

    slide(doc, 4, "Arquitectura de la solución", "0:50",
        "El diagrama de capas.",
        "La solución tiene cuatro capas. Toda entrada pasa primero por una capa de seguridad; "
        "luego el agente principal, con LangChain y el ciclo ReAct, decide qué herramienta o "
        "sub-agente usar, apoyándose en RAG y en la memoria; y cada paso queda instrumentado por "
        "la capa de observabilidad. Integra así RA1, RA2 y RA3.",
        "“Veámoslo funcionando.”")

    slide(doc, 5, "Agente funcional", "1:15",
        "Las 3 columnas: herramientas, memoria y planificación.",
        "El agente integra cinco herramientas de consulta, escritura y razonamiento; memoria de "
        "corto plazo con una ventana de cinco turnos y de largo plazo en disco, que da continuidad "
        "entre sesiones; y planificación Plan-and-Execute con orquestación CrewAI de cuatro roles "
        "especializados.",
        "“Y esto no es teoría: aquí está la evidencia de que funciona.”")

    slide(doc, 6, "Evidencia de ejecución — pipeline y pruebas", "0:45",
        "Capturas de consola: pipeline con KPIs y las 22 pruebas en verde.",
        "Esta es la ejecución real del sistema. El pipeline procesa 688 interacciones y reporta "
        "los indicadores: 85% de precisión, la latencia y el uso de recursos. Y abajo, las 22 "
        "pruebas automatizadas pasan en verde, validando seguridad, observabilidad y RAG.",
        "“También dejamos evidencia de la seguridad y del RAG.”")

    slide(doc, 7, "Evidencia de ejecución — seguridad y RAG", "0:45",
        "Captura de consola: guardrails y recuperación RAG.",
        "Aquí se ve la seguridad en acción: el RUT y el correo quedan enmascarados y un intento de "
        "prompt-injection es bloqueado. Y el RAG recupera combinando fuentes internas del banco "
        "con fuentes externas, como la CMF, el SERNAC y la Ley 19.628.",
        "“Con el agente demostrado, veamos el diseño que hay detrás.”")

    slide(doc, 8, "Diseño LLM + RAG", "0:40",
        "Viñetas y el recuadro de recuperación combinada.",
        "El comportamiento se controla con prompts —rol, reglas, few-shot y chain-of-thought— y "
        "las respuestas sobre productos se fundamentan con RAG, combinando fuentes internas del "
        "banco con fuentes externas de referencia normativa.",
        "“¿Y cómo se comporta el agente en el tiempo? Lo medimos.”")

    slide(doc, 9, "Observabilidad — resultados", "1:05",
        "Las tarjetas de KPI y los gráficos de precisión y latencia.",
        "Instrumentamos cada interacción. Sobre 688 interacciones, el agente logra 85% de "
        "precisión, 6,9% de error y una latencia p95 de 8,3 segundos. Por escenario se ve que las "
        "consultas claras superan el 90% y que el multi-agente domina la latencia.",
        "“Esos números además nos revelaron problemas concretos.”")

    slide(doc, 10, "Trazabilidad — hallazgos clave", "0:55",
        "El gráfico temporal con el incidente y las viñetas.",
        "El análisis de trazabilidad detectó automáticamente un incidente el 11 de junio, con "
        "latencia y error tres desviaciones sobre lo normal. Identificó el cuello de botella en "
        "la orquestación multi-agente y que el principal predictor de éxito es elegir bien la "
        "herramienta: las entradas ambiguas hunden la precisión 37 puntos.",
        "“Todo esto bajo una capa de seguridad y uso responsable.”")

    slide(doc, 11, "Seguridad y uso responsable", "0:45",
        "Los cinco controles y el recuadro normativo.",
        "La seguridad es transversal: enmascaramos datos personales antes de cualquier registro, "
        "bloqueamos prompt-injection, limitamos la tasa de uso, auditamos con integridad y "
        "exigimos confirmación humana en cada transferencia. Todo alineado con la Ley 19.628, la "
        "CMF y OWASP. [Añade aquí tu reflexión ética propia.]",
        "“Con esta evidencia, propusimos mejoras priorizadas.”")

    slide(doc, 12, "Mejoras propuestas", "0:40",
        "Las cinco mejoras numeradas.",
        "Cada mejora responde a un hallazgo: desambiguar la intención para subir la precisión; "
        "paralelizar el multi-agente para bajar latencia y costo; añadir resiliencia ante "
        "incidentes; reforzar el guardrail de seguridad; y activar alertas proactivas. Priorizamos "
        "la desambiguación por su impacto.",
        "“Para cerrar.”")

    slide(doc, 13, "Cierre", "0:15",
        "Diapositiva de cierre.",
        "En síntesis, entregamos una solución medible, segura y escalable, que integra LLM, RAG, "
        "agentes y observabilidad de extremo a extremo. Quedo atento a sus preguntas. Muchas "
        "gracias.",
        None)

    h1(doc, "Preguntas frecuentes (prepara respuestas propias)")
    vinneta(doc, [
        [("¿Cómo medirías si una mejora funcionó? ", True),
         ("Comparando las métricas antes/después sobre el mismo set de escenarios.", False)],
        [("¿Qué pasa si el modelo se cae? ", True),
         ("Timeouts, reintentos con backoff y un modelo de respaldo (fallback).", False)],
        [("¿Cómo evitas filtrar datos personales? ", True),
         ("Enmascaramiento de PII antes de logs/prompts y auditoría con hash.", False)],
        [("¿Por qué LangChain y CrewAI? ", True),
         ("LangChain por el ciclo ReAct y la gestión de memoria/herramientas; CrewAI por roles "
          "especializados y replanificación jerárquica.", False)],
        [("¿Cómo escalas a 10× el tráfico? ", True),
         ("Cache de productos, paralelización de especialistas, colas y autoescalado por demanda.", False)],
        [("¿Cuál fue el mayor desafío? ", True),
         ("[Responde con tu experiencia real del proyecto.]", False)],
    ], size=10)

    caja(doc, "Recordatorios finales",
        "• La evidencia de ejecución va en las diapositivas 6 y 7 (capturas reales), así no "
        "dependes de una demo en vivo. Si igual quieres mostrar el sistema, ten el notebook listo.\n"
        "• Habla con lenguaje técnico y respalda con cifras (IE13–IE15).\n"
        "• La claridad y la autocrítica puntúan más que la perfección técnica.")

    doc.save(SALIDA)
    return SALIDA


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    print(f"✔ Guion generado: {os.path.relpath(build(), ROOT)}")
