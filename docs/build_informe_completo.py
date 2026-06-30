"""
Informe Integral del Proyecto (Word) — consolida EP3 + EFT en un solo documento
================================================================================
Documento completo de extremo a extremo del agente "Banco Digital Chile":
caso organizacional, diseño LLM+RAG, agente funcional, observabilidad,
trazabilidad, seguridad, resultados, pruebas de software y conclusiones.

Lee las métricas reales del pipeline para mantener coherencia con los artefactos.
Uso:
    python run_observability.py          # genera datos/gráficos
    python docs/build_informe_completo.py
"""

from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx.shared import Pt, Inches, RGBColor

from docs._docx_helpers import (
    ROOT, nuevo_documento, bloque_titulo, set_page_numbers,
    h1, h2, h3, parrafo, rico, vinneta, numerada, tabla, figura, caja, salto_pagina,
)
from observability import metrics as M
from observability.analyze_traces import patrones, anomalias_temporales, puntos_de_falla

AUTORES = "Roger Rosas · Rodrigo Santis · Edgardo Gutiérrez"
FECHA = "Junio 2025"
ENTREGA = os.path.join(ROOT, "entrega"); os.makedirs(ENTREGA, exist_ok=True)
SALIDA = os.path.join(ENTREGA, "Informe_Completo_BancoDigitalChile.docx")
GRIS = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xC0, 0x39, 0x2B)

NOMBRES = {
    "consulta_saldo": "Consulta de saldo", "info_producto": "Info. de producto",
    "simulacion_credito": "Simulación crédito", "transferencia": "Transferencia",
    "registro_reclamo": "Registro reclamo", "multiagente_credito": "Multi-agente crédito",
    "consulta_ambigua": "Consulta ambigua", "seguridad_injection": "Intento injection",
}


def pct(x): return f"{x*100:.1f}%"


def codigo(doc, lineas):
    for ln in lineas:
        p = doc.add_paragraph()
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.left_indent = Inches(0.2)


def nota_ia(doc, texto):
    t = doc.add_paragraph()
    r = t.add_run("⚠ " + texto)
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = ROJO
    t.paragraph_format.space_after = Pt(6)


def andamiaje(doc, guias):
    for g in guias:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(g); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRIS
        p.paragraph_format.space_after = Pt(2)
    ph = doc.add_paragraph()
    rr = ph.add_run("[ Espacio para redacción propia del equipo / estudiante ]")
    rr.font.size = Pt(9); rr.italic = True; rr.font.color.rgb = GRIS


def _rag_ejemplo():
    from rag.rag_pipeline import responder_con_rag
    return responder_con_rag(
        "¿Qué tasa tiene el crédito de consumo y qué información debe entregarme el banco?", k=3)


def build():
    df = M.load_df()
    g = M.global_metrics(df)
    esc = M.por_escenario(df)
    cons_tabla, cons = M.consistencia(df)
    err = M.errores(df)
    sec = M.seguridad(df)
    serie = M.serie_diaria(df)
    anom = anomalias_temporales(serie)
    pat = patrones(df)
    fallas = puntos_de_falla(df)

    doc = nuevo_documento()
    set_page_numbers(doc, "Informe Integral · Banco Digital Chile · ISY0101")

    bloque_titulo(
        doc,
        titulo="Informe Integral del Proyecto",
        subtitulo="Agente conversacional “Banco Digital Chile” — Solución de IA para banca digital",
        sigla="ISY0101 · Ingeniería de Soluciones con Inteligencia Artificial",
        autores=AUTORES, fecha=FECHA,
        ponderacion="Consolidado EP3 (Observabilidad) + EFT (Final Transversal)",
    )

    # Resumen ejecutivo
    h1(doc, "Resumen ejecutivo")
    rico(doc, [
        ("Este informe integra de extremo a extremo el proyecto desarrollado durante el "
         "semestre: el diseño de una solución basada en LLM y RAG, la construcción de un agente "
         "funcional y la incorporación de observabilidad, trazabilidad y seguridad de nivel "
         "productivo para ", False),
        ("Banco Digital Chile", True),
        (". El agente automatiza tareas de banca digital (consulta de saldo, transferencias, "
         "simulación de créditos, reclamos e información de productos) usando LangChain (ReAct), "
         "un pipeline RAG sobre productos y normativa, memoria de corto y largo plazo, "
         "planificación multi-etapa y orquestación multi-agente con CrewAI. Sobre ", False),
        (f"{g['interacciones_totales']} interacciones monitoreadas", True),
        (f", alcanza una precisión de {pct(g['precision_exito'])}, latencia media de "
         f"{g['latencia_media_ms']:.0f} ms y una tasa de error de {pct(g['tasa_error'])}, con "
         "hallazgos de mejora priorizados y una capa de seguridad alineada a la Ley 19.628, la "
         "CMF y OWASP.", False),
    ])

    h2(doc, "Contenido")
    vinneta(doc, [
        "1. Análisis del caso organizacional",
        "2. Diseño de la solución basada en LLM y RAG",
        "3. Desarrollo del agente funcional",
        "4. Observabilidad y trazabilidad",
        "5. Seguridad y uso responsable",
        "6. Recomendaciones de optimización",
        "7. Resultados consolidados y pruebas de software",
        "8. Declaración de uso de IA, conclusiones y referencias",
    ], size=10)
    salto_pagina(doc)

    # ── 1. Caso ───────────────────────────────────────────────────────────────
    h1(doc, "1. Análisis del caso organizacional")
    h2(doc, "1.1 Contexto y problema")
    parrafo(doc,
        "Banco Digital Chile es una entidad de operación 100% digital cuyo canal de atención "
        "recibe un alto volumen de consultas repetitivas (saldos, condiciones de productos, "
        "simulaciones) junto con operaciones sensibles (transferencias) y reclamos. La atención "
        "humana no escala a costos razonables y genera tiempos de espera elevados; un chatbot de "
        "reglas no comprende lenguaje natural ni ejecuta acciones de extremo a extremo.")
    h2(doc, "1.2 Requerimientos y desafíos")
    vinneta(doc, [
        [("Funcionales: ", True), ("resolver consultas y ejecutar operaciones manteniendo el "
         "contexto de la conversación y del cliente.", False)],
        [("No funcionales: ", True), ("baja latencia, trazabilidad de cada decisión, seguridad "
         "y cumplimiento normativo (Ley 19.628, CMF), y escalabilidad ante demanda variable.", False)],
        [("Desafíos: ", True), ("ambigüedad del lenguaje, control de alucinaciones, protección "
         "de datos personales y resiliencia ante fallos del proveedor del modelo.", False)],
    ])

    # ── 2. Diseño ─────────────────────────────────────────────────────────────
    h1(doc, "2. Diseño de la solución basada en LLM y RAG")
    h2(doc, "2.1 Formulación de prompts")
    parrafo(doc,
        "El agente usa un system prompt que fija rol, alcance y reglas de seguridad, combinado "
        "con few-shot implícito en las descripciones de herramientas y razonamiento "
        "chain-of-thought para tareas multi-etapa (técnicas desarrolladas en RA1/IL1.2). "
        "Extracto del prompt de sistema:")
    codigo(doc, [
        "Eres un asistente bancario inteligente de Banco Digital Chile.",
        "Reglas:",
        "1. Para mostrar saldo/datos sensibles, solicita el RUT si no lo tienes.",
        "2. Para transferencias, confirma origen, destino y monto antes de ejecutar.",
        "3. Para reclamos, solicita una descripción clara del problema.",
        "4. Ante solicitudes complejas, planifica los pasos antes de actuar.",
    ])

    h2(doc, "2.2 Pipeline RAG (fuentes internas + externas)")
    parrafo(doc,
        "La consulta de productos y políticas se resuelve con recuperación aumentada "
        "(rag/rag_pipeline.py): el conocimiento se divide en fragmentos (chunking), se vectoriza "
        "con TF-IDF y se recupera por similitud coseno para enriquecer la respuesta del modelo. "
        "La base combina dos tipos de fuente:")
    vinneta(doc, [
        [("Fuentes internas: ", True), ("catálogo de productos (tasas, plazos, requisitos) y "
         "políticas operativas del banco (transferencias, reclamos).", False)],
        [("Fuentes externas: ", True), ("marco normativo y de protección al consumidor "
         "(CMF, SERNAC, Ley 19.628).", False)],
    ])
    rag = _rag_ejemplo()
    parrafo(doc,
        f"Ejemplo real de recuperación para «{rag['pregunta']}» (combina ambos orígenes: "
        f"{', '.join(rag['origenes_combinados'])}):", space_after=4)
    tabla(doc, ["Origen", "Fuente", "Score"],
          [[f["origen"], f["fuente"].split(" — ")[0].split(" (")[0], f["score"]]
           for f in rag["fuentes"]], anchos=[1.1, 4.4, 0.9], size=8.5)
    parrafo(doc,
        "El backend léxico (TF-IDF) garantiza reproducibilidad sin credenciales; en producción "
        "se sustituye por embeddings densos manteniendo la interfaz.", space_after=4)

    h2(doc, "2.3 Arquitectura de la solución")
    parrafo(doc,
        "La arquitectura integra el LLM, la recuperación de información y el control de contexto "
        "(memoria) bajo una capa de seguridad de entrada y una capa de observabilidad de salida. "
        "Cada solicitud atraviesa el saneamiento de seguridad, es procesada por el AgentExecutor "
        "(ciclo ReAct) que decide qué herramienta o sub-agente invocar, y queda instrumentada.")
    figura(doc, "00_arquitectura.png",
           "Figura 1. Arquitectura por capas, integrando RA1 (LLM/RAG), RA2 (agente) y RA3 (observabilidad).",
           width=6.2)

    h2(doc, "2.4 Justificación de decisiones de diseño")
    nota_ia(doc, "Justificación técnica: el equipo debe redactarla con sus propias palabras "
                 "(política de uso de IA). Andamiaje de apoyo:")
    andamiaje(doc, [
        "¿Por qué LangChain para el agente principal? (ReAct, gestión de herramientas, memoria).",
        "¿Por qué CrewAI para la orquestación? (roles, delegación, replanificación jerárquica).",
        "¿Por qué ventana k=5 y JSON en memoria? (equilibrio contexto/costo; simplicidad).",
        "¿Por qué gpt-4o vía GitHub Models? (capacidad, disponibilidad, costo).",
        "Relación de cada decisión con requerimientos y limitaciones del modelo.",
    ])

    # ── 3. Agente ─────────────────────────────────────────────────────────────
    h1(doc, "3. Desarrollo del agente funcional")
    h2(doc, "3.1 Herramientas de consulta, escritura y razonamiento")
    tabla(doc, ["Herramienta", "Tipo", "Función"], [
        ["consultar_saldo", "Consulta", "Saldo, tipo de cuenta y productos por RUT"],
        ["realizar_transferencia", "Escritura", "Transferencia validando saldo suficiente"],
        ["calcular_cuota_credito", "Razonamiento", "Amortización francesa (cuota, intereses)"],
        ["registrar_reclamo", "Escritura", "Ticket de reclamo persistido con SLA"],
        ["buscar_producto_bancario", "Consulta/RAG", "Características y requisitos de productos"],
    ], anchos=[2.0, 1.3, 4.0], size=9, fuente_resaltar_col0=True)

    h2(doc, "3.2 Memoria y recuperación de contexto")
    parrafo(doc,
        "Memoria de corto plazo con ConversationBufferWindowMemory (k=5) para sostener el hilo "
        "de la conversación, y memoria de largo plazo en JSON persistente por RUT que inyecta el "
        "historial del cliente en el system prompt, asegurando continuidad entre sesiones.")
    h2(doc, "3.3 Planificación y orquestación")
    parrafo(doc,
        "Para solicitudes multi-etapa se aplica Plan-and-Execute (el agente genera un plan antes "
        "de actuar). Las solicitudes complejas se derivan a CrewAI con cuatro roles —Gerente, "
        "Analista Financiero, Asesor de Productos y Especialista en Operaciones— en procesos "
        "secuencial y jerárquico (el Gerente delega y replanifica dinámicamente y consolida la "
        "respuesta final al cliente).")
    h2(doc, "3.4 Documentación del flujo automatizado")
    parrafo(doc,
        "El flujo completo —entrada del cliente → seguridad → agente (ReAct) → herramientas/RAG/"
        "memoria/orquestación → instrumentación → respuesta— está documentado en el repositorio "
        "(README_EP2 y README_EP3), con escenarios de prueba y diagramas de secuencia.")

    # ── 4. Observabilidad ─────────────────────────────────────────────────────
    h1(doc, "4. Observabilidad y trazabilidad")
    h2(doc, "4.1 Métricas de observabilidad implementadas")
    parrafo(doc,
        "Cada interacción se instrumenta con un tracer que emite trazas JSONL "
        "(logs/agent_traces.jsonl) más un log legible (logs/agent.log), con PII enmascarada. Se "
        "implementaron cinco métricas en dos dimensiones:")
    vinneta(doc, [
        [("Precisión: ", True), ("tasa de éxito y acierto de herramienta vs. verdad de referencia.", False)],
        [("Consistencia: ", True), ("estabilidad de la decisión ante entradas equivalentes.", False)],
        [("Frecuencia de errores: ", True), ("proporción y tipo de fallos.", False)],
        [("Latencia: ", True), ("media, p50, p95 y máx.", False)],
        [("Uso de recursos: ", True), ("tokens, costo estimado y pasos por tarea.", False)],
    ])
    caja(doc, f"KPIs globales (14 días · {g['interacciones_totales']} interacciones)",
        f"Precisión: {pct(g['precision_exito'])}   ·   Acierto herramienta: {pct(g['acierto_herramienta'])}   ·   "
        f"Consistencia: {pct(cons['consistencia_media'])}   ·   Error: {pct(g['tasa_error'])}\n"
        f"Latencia media/p95/máx: {g['latencia_media_ms']:.0f} / {g['latencia_p95_ms']:.0f} / "
        f"{g['latencia_max_ms']:.0f} ms   ·   Tokens: {g['tokens_totales']:,}   ·   "
        f"Costo: US${g['costo_total_usd']:.2f}")

    h2(doc, "4.2 Análisis de registros y trazabilidad")
    falla_top = fallas.iloc[0]
    cuello = esc.index[0]
    items = [
        [("Punto de falla dominante: ", True),
         (f"«{NOMBRES.get(falla_top['escenario'], falla_top['escenario'])}» por "
          f"{falla_top['motivo'].replace('_',' ')} ({falla_top['pct_de_fallos']:.0f}% de los fallos).", False)],
        [("Cuello de botella: ", True),
         (f"«{NOMBRES.get(cuello, cuello)}» con p95 de {esc.iloc[0]['latencia_p95_ms']:.0f} ms "
          "(orquestación multi-agente).", False)],
        [("Predictor de éxito: ", True),
         (f"herramienta correcta → éxito {pct(pat['exito_con_tool_correcta'])}; incorrecta → "
          f"{pct(pat['exito_con_tool_incorrecta'])}.", False)],
        [("Correlación recursos–latencia: ", True),
         (f"r = {pat['corr_tokens_latencia']} entre tokens y latencia.", False)],
    ]
    vinneta(doc, items)
    figura(doc, "02_latencia_por_escenario.png",
           "Figura 2. Latencia p95 vs. media por escenario.", width=5.5)

    h2(doc, "4.3 Patrones y anomalías")
    dia = anom[anom["anomalia"]]
    if not dia.empty:
        d = dia.iloc[0]
        parrafo(doc,
            f"Se detectó automáticamente un incidente el {d['fecha']}: latencia p95 = "
            f"{d['latencia_p95_ms']:.0f} ms (z = {d['z_latencia_p95_ms']}) y tasa de error = "
            f"{pct(d['tasa_error'])} (z = {d['z_tasa_error']}), compatible con una degradación del "
            "proveedor del modelo. Además, la precisión cae de "
            f"{pct(pat['precision_resto'])} (entradas claras) a {pct(pat['precision_ambigua'])} "
            f"(ambiguas), una brecha de {pat['brecha_ambiguedad']*100:.0f} pp.")
    figura(doc, "03_serie_temporal_incidente.png",
           "Figura 3. Evolución diaria de latencia p95 y tasa de error (incidente sombreado).",
           width=5.8)

    h2(doc, "4.4 Dashboard de monitoreo")
    rico(doc, [
        ("El dashboard se construyó en ", False), ("Power BI", True),
        (" sobre la tabla de hechos dashboard/powerbi_dataset.csv con medidas DAX "
         "(guía en dashboard/GUIA_POWERBI.md), e incluye una vista rápida en Excel. Captura del "
         "tablero integrado:", False),
    ])
    figura(doc, "07_dashboard_panel.png",
           "Figura 4. Dashboard de monitoreo integrado (KPIs + visualizaciones).", width=6.4)

    # ── 5. Seguridad ──────────────────────────────────────────────────────────
    h1(doc, "5. Seguridad y uso responsable")
    parrafo(doc,
        "La capa transversal de seguridad (security/guardrails.py) está alineada con la Ley "
        "19.628 sobre datos personales, la normativa de ciberseguridad de la CMF y el OWASP Top "
        "10 for LLM Applications:")
    vinneta(doc, [
        [("Privacidad por diseño: ", True), ("enmascaramiento de PII (RUT, tarjeta, correo, "
         "teléfono) antes de logs y prompts; los RUT conservan solo 3 dígitos.", False)],
        [("Defensa ante prompt-injection: ", True), ("bloqueo de intentos de ignorar "
         "instrucciones, revelar el system prompt u operar sin confirmación (OWASP LLM01).", False)],
        [("Rate limiting y validación de entrada: ", True), ("control de abuso y saneamiento.", False)],
        [("Auditoría con integridad: ", True), ("bitácora con encadenamiento de hash (no-repudio).", False)],
        [("Confirmación humana: ", True), ("toda transferencia requiere validación explícita.", False)],
    ])
    rico(doc, [
        (f"Sobre la telemetría se registraron {sec['intentos_injection']} intentos de "
         f"prompt-injection, bloqueados en {pct(sec['tasa_bloqueo'])}", True),
        (f"; persisten {sec['bypasses']} casos no bloqueados (riesgo residual a reforzar). Se "
         f"enmascaró PII en {sec['interacciones_con_pii']} interacciones. Las consideraciones "
         "éticas (transparencia, no maleficencia, equidad, supervisión humana) se detallan en "
         "security/policies.md.", False),
    ])

    # ── 6. Recomendaciones ────────────────────────────────────────────────────
    h1(doc, "6. Recomendaciones de optimización")
    parrafo(doc, "Priorizadas por impacto en sostenibilidad y escalabilidad, fundamentadas en "
                 "las métricas y la trazabilidad:")
    numerada(doc, [
        [("Desambiguación de intención. ", True),
         (f"La brecha de {pat['brecha_ambiguedad']*100:.0f} pp en entradas ambiguas justifica una "
          "capa de clarificación (repreguntar) y few-shot de enrutamiento. Mayor impacto en precisión.", False)],
        [("Optimización del cuello de botella multi-agente. ", True),
         ("Paralelizar especialistas, cachear productos y usar un modelo liviano para sub-tareas "
          "reduce latencia p95 y consumo de tokens.", False)],
        [("Resiliencia ante incidentes. ", True),
         ("Reintentos con backoff, timeouts y modelo de respaldo (fallback) mitigan episodios "
          "como el incidente detectado.", False)],
        [("Refuerzo del guardrail de seguridad. ", True),
         (f"Los {sec['bypasses']} bypasses exigen un clasificador LLM dedicado y canary tokens "
          "para acercar el bloqueo al 100%.", False)],
        [("Monitoreo proactivo con alertas. ", True),
         ("Umbrales (p95 > histórico + 2σ, error > 10%) que disparen alertas sobre el dashboard.", False)],
    ])

    # ── 7. Resultados + pruebas ───────────────────────────────────────────────
    h1(doc, "7. Resultados consolidados y pruebas de software")
    h2(doc, "7.1 Resultados por escenario")
    filas = [[NOMBRES.get(n, n), int(r["n"]), pct(r["precision_exito"]),
              pct(r["tasa_error"]), f"{r['latencia_p95_ms']:.0f}", f"{int(r['tokens_medio'])}"]
             for n, r in esc.iterrows()]
    tabla(doc, ["Escenario", "N", "Precisión", "Tasa error", "Latencia p95 (ms)", "Tokens medio"],
          filas, anchos=[1.9, 0.6, 1.0, 1.0, 1.5, 1.1], size=9, fuente_resaltar_col0=True)

    h2(doc, "7.2 Evidencia de pruebas de software")
    parrafo(doc,
        "El proyecto incluye una batería de 22 pruebas automatizadas (pytest) que validan la "
        "seguridad (enmascaramiento, injection, rate-limit), la observabilidad (instrumentación, "
        "métricas, anomalías) y el RAG (recuperación y combinación de fuentes). Todas pasan; la "
        "salida se archiva en reports/evidencia_pruebas.txt.")
    tabla(doc, ["Suite", "Cobertura", "Resultado"], [
        ["test_guardrails.py", "PII, injection, rate-limit, auditoría", "9/9 ✓"],
        ["test_observability.py", "tracer, métricas, consistencia, anomalías", "8/8 ✓"],
        ["test_rag.py", "recuperación y combinación de fuentes", "5/5 ✓"],
    ], anchos=[2.2, 3.8, 1.2], size=9, fuente_resaltar_col0=True)

    # ── 8. Declaración + conclusiones + refs ──────────────────────────────────
    h1(doc, "8. Declaración de uso de IA, conclusiones y referencias")
    h2(doc, "8.1 Declaración de uso de inteligencia artificial")
    parrafo(doc,
        "Conforme a la guía de Biblioteca Duoc UC (bibliotecas.duoc.cl/ia), se declara el uso de "
        "IA como apoyo en redacción y estructura del informe, generación de diagramas y "
        "organización de tablas, y en la generación de la instrumentación y telemetría sintética "
        "(revisada por el equipo).")
    nota_ia(doc, "Las conclusiones, justificaciones técnicas y reflexiones individuales fueron "
                 "redactadas por los integrantes sin apoyo de IA.")

    h2(doc, "8.2 Conclusiones del equipo")
    nota_ia(doc, "Completar sin IA.")
    andamiaje(doc, [
        "Grado de cumplimiento de los requerimientos del caso.",
        "Principales aprendizajes de integrar LLM + RAG + agentes + observabilidad.",
        "Limitaciones y trabajo futuro.",
    ])
    for integrante in ["Roger Rosas", "Rodrigo Santis", "Edgardo Gutiérrez"]:
        h3(doc, f"Reflexión individual — {integrante}")
        andamiaje(doc, [
            "Mi contribución específica al proyecto.",
            "Lo que aprendí y cómo lo aplicaría en un contexto real.",
            "Un dilema ético o de seguridad que identifiqué y cómo lo abordaría.",
        ])

    h2(doc, "8.3 Referencias")
    refs = [
        "Biblioteca Duoc UC. (2024). Guía para el uso de inteligencia artificial. https://bibliotecas.duoc.cl/ia",
        "Chase, H. (2023). LangChain: Building applications with LLMs through composability. https://www.langchain.com/",
        "Comisión para el Mercado Financiero. (2020). RAN Cap. 20-7: Gestión de la ciberseguridad. CMF Chile.",
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. arXiv:2005.11401.",
        "Ley N°19.628. Sobre protección de la vida privada. Biblioteca del Congreso Nacional de Chile.",
        "Moura, J. (2024). CrewAI: Orchestrating role-playing autonomous AI agents. https://docs.crewai.com/",
        "OWASP Foundation. (2023). OWASP Top 10 for Large Language Model Applications.",
        "Yao, S., et al. (2022). ReAct: Synergizing reasoning and acting in language models. arXiv:2210.03629.",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.runs[0].font.size = Pt(9); p.paragraph_format.space_after = Pt(3)

    parrafo(doc,
        "Repositorio (código, RAG, notebook ejecutable, README, dashboard y evidencia de "
        "pruebas): github.com/rogrosas/Ingenier-a-de-Soluciones-con-Inteligencia-Artificial. "
        "Reproducible con «python run_observability.py».", size=9, italic=True)

    # ── Anexo A: Evidencias de ejecución ──────────────────────────────────────
    salto_pagina(doc)
    h1(doc, "Anexo A — Evidencias de ejecución (capturas)")
    parrafo(doc,
        "Capturas de la ejecución real de los componentes, reproducibles ejecutando "
        "«entrega/evidencia/generar_evidencia.bat» o «python run_observability.py». La salida "
        "completa se archiva en «entrega/evidencia/salida_ejecucion.txt».")
    figura(doc, "ev_01_pipeline.png",
           "Evidencia 1. Ejecución del pipeline de observabilidad: KPIs globales sobre 688 "
           "interacciones (precisión 85,0%, latencia y uso de recursos).", width=6.3)
    figura(doc, "ev_02_pruebas.png",
           "Evidencia 2. Suite de pruebas automatizadas: 22/22 pruebas PASSED (pytest).", width=6.3)
    figura(doc, "ev_03_seg_rag.png",
           "Evidencia 3. Seguridad (enmascaramiento de PII y bloqueo de prompt-injection) y "
           "RAG (recuperación combinando fuentes internas y externas).", width=6.3)
    figura(doc, "07_dashboard_panel.png",
           "Evidencia 4. Dashboard de monitoreo integrado (KPIs + visualizaciones).", width=6.3)

    doc.save(SALIDA)
    return SALIDA


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    print(f"✔ Informe completo generado: {os.path.relpath(build(), ROOT)}")
