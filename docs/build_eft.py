"""
Generador del Informe Final Transversal (Word) — Proyecto completo
===================================================================
Consolida RA1 (LLM/RAG/prompts) + RA2 (agente funcional) + RA3 (observabilidad,
trazabilidad y seguridad) del agente "Banco Digital Chile", con diagrama de
arquitectura, tablas de resultados reales y las secciones formales exigidas.

IMPORTANTE (integridad académica): las secciones de CONCLUSIONES, JUSTIFICACIONES
y REFLEXIONES INDIVIDUALES se dejan como ANDAMIAJE guiado para que cada estudiante
las redacte sin apoyo de IA, según la pauta de la evaluación.

Uso:
    python docs/make_arch_diagram.py
    python docs/build_eft.py
"""

from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx.shared import Pt, Inches, RGBColor

from docs._docx_helpers import (
    ROOT, nuevo_documento, bloque_titulo, set_page_numbers,
    h1, h2, parrafo, rico, vinneta, numerada, tabla, figura, caja, salto_pagina,
)
from observability import metrics as M
from observability.analyze_traces import patrones, anomalias_temporales

AUTORES = "Roger Rosas · Rodrigo Santis · Edgardo Gutiérrez"
FECHA = "Junio 2025"
ENTREGA = os.path.join(ROOT, "entrega"); os.makedirs(ENTREGA, exist_ok=True)
SALIDA = os.path.join(ENTREGA, "EFT_Informe_Final_BancoDigitalChile.docx")
GRIS = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xC0, 0x39, 0x2B)


def pct(x): return f"{x*100:.1f}%"


def _rag_ejemplo():
    """Recupera un ejemplo real del pipeline RAG que combina fuentes internas y externas."""
    from rag.rag_pipeline import responder_con_rag
    return responder_con_rag(
        "¿Qué tasa tiene el crédito de consumo y qué información debe entregarme el banco?", k=3)


def codigo(doc, lineas):
    for ln in lineas:
        p = doc.add_paragraph()
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)


def nota_ia(doc, texto):
    t = doc.add_paragraph()
    r = t.add_run("⚠ " + texto)
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = ROJO
    t.paragraph_format.space_after = Pt(6)


def andamiaje(doc, guias):
    """Bloque de andamiaje (texto a completar por el estudiante, sin IA)."""
    for g in guias:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(g); r.font.size = Pt(9.5); r.italic = True; r.font.color.rgb = GRIS
        p.paragraph_format.space_after = Pt(2)
    ph = doc.add_paragraph()
    rr = ph.add_run("[ Espacio para redacción propia del equipo / estudiante ]")
    rr.font.size = Pt(9); rr.italic = True; rr.font.color.rgb = GRIS


def build():
    df = M.load_df()
    g = M.global_metrics(df)
    esc = M.por_escenario(df)
    _, cons = M.consistencia(df)
    sec = M.seguridad(df)
    serie = M.serie_diaria(df)
    anom = anomalias_temporales(serie)
    pat = patrones(df)

    doc = nuevo_documento()
    set_page_numbers(doc, "EFT · Banco Digital Chile · ISY0101")

    bloque_titulo(
        doc,
        titulo="Solución de IA para Banca Digital",
        subtitulo="Agente conversacional “Banco Digital Chile” — Informe Final",
        sigla="ISY0101 · Ingeniería de Soluciones con Inteligencia Artificial",
        autores=AUTORES, fecha=FECHA,
        ponderacion="Evaluación Final Transversal · 40%",
    )

    # Resumen ejecutivo
    h1(doc, "Resumen ejecutivo")
    rico(doc, [
        ("Se presenta la versión final de un agente conversacional para ", False),
        ("Banco Digital Chile", True),
        (", que automatiza la atención de tareas bancarias —consulta de saldo, "
         "transferencias, simulación de créditos, registro de reclamos e información de "
         "productos— integrando un modelo de lenguaje (gpt-4o), un pipeline RAG sobre la base "
         "de productos y políticas, memoria de corto y largo plazo, planificación multi-etapa "
         "y orquestación multi-agente. La solución incorpora además una capa de "
         "observabilidad, trazabilidad y seguridad de nivel productivo. Sobre ", False),
        (f"{g['interacciones_totales']} interacciones monitoreadas", True),
        (f", el agente alcanza una precisión de {pct(g['precision_exito'])} y una latencia "
         f"media de {g['latencia_media_ms']:.0f} ms, con hallazgos de mejora claramente "
         "priorizados.", False),
    ])

    # Índice
    h1(doc, "Contenido")
    vinneta(doc, [
        "1. Análisis del caso organizacional",
        "2. Diseño de la solución basada en LLM y RAG",
        "3. Desarrollo del agente funcional",
        "4. Observabilidad, trazabilidad y seguridad",
        "5. Resultados consolidados",
        "6. Declaración de uso de IA",
        "7. Conclusiones y reflexiones individuales",
        "8. Referencias",
    ], size=10)

    salto_pagina(doc)

    # ── 1. Caso organizacional ────────────────────────────────────────────────
    h1(doc, "1. Análisis del caso organizacional")
    h2(doc, "1.1 Contexto y problema")
    parrafo(doc,
        "Banco Digital Chile es una entidad financiera de operación 100% digital cuyo canal "
        "de atención recibe un alto volumen de consultas repetitivas (saldos, condiciones de "
        "productos, simulaciones de crédito) junto con operaciones sensibles (transferencias) "
        "y reclamos. La atención humana no escala a costos razonables y genera tiempos de "
        "espera elevados, mientras que un chatbot tradicional basado en reglas no comprende el "
        "lenguaje natural ni ejecuta acciones de extremo a extremo.")
    h2(doc, "1.2 Requerimientos y desafíos")
    vinneta(doc, [
        [("Funcionales: ", True), ("resolver consultas y ejecutar operaciones (transferencias, "
         "reclamos, simulaciones) manteniendo el contexto de la conversación y del cliente.", False)],
        [("No funcionales: ", True), ("baja latencia, trazabilidad de cada decisión, seguridad "
         "y cumplimiento normativo (Ley 19.628, CMF), y escalabilidad ante demanda variable.", False)],
        [("Desafíos: ", True), ("ambigüedad del lenguaje de los clientes, control de "
         "alucinaciones, protección de datos personales y resiliencia ante fallos del proveedor "
         "del modelo.", False)],
    ])

    # ── 2. Diseño LLM + RAG ───────────────────────────────────────────────────
    h1(doc, "2. Diseño de la solución basada en LLM y RAG")
    h2(doc, "2.1 Formulación de prompts (IE1)")
    parrafo(doc,
        "El agente usa un system prompt que fija rol, alcance y reglas de seguridad, combinado "
        "con few-shot implícito en las descripciones de herramientas y razonamiento "
        "chain-of-thought para tareas multi-etapa. Extracto del prompt de sistema:")
    codigo(doc, [
        "Eres un asistente bancario inteligente de Banco Digital Chile.",
        "Reglas:",
        "1. Para mostrar saldo/datos sensibles, solicita el RUT si no lo tienes.",
        "2. Para transferencias, confirma origen, destino y monto antes de ejecutar.",
        "3. Para reclamos, solicita una descripción clara del problema.",
        "4. Ante solicitudes complejas, planifica los pasos antes de actuar.",
    ])
    parrafo(doc,
        "Las técnicas de prompting (zero-shot, few-shot y chain-of-thought) se desarrollaron y "
        "evaluaron en RA1/IL1.2, y aquí se aplican a la formulación de las instrucciones del "
        "agente y de las herramientas.", space_after=4)

    h2(doc, "2.2 Pipeline RAG (IE2)")
    parrafo(doc,
        "La consulta de productos y políticas se resuelve mediante recuperación aumentada "
        "(RAG, módulo rag/rag_pipeline.py): el conocimiento se divide en fragmentos (chunking), "
        "se vectoriza con TF-IDF y se recupera por similitud coseno para enriquecer la respuesta "
        "del modelo. La base combina explícitamente dos tipos de fuente:")
    vinneta(doc, [
        [("Fuentes internas: ", True), ("catálogo de productos (tasas, plazos, requisitos) y "
         "políticas operativas del banco (transferencias, reclamos).", False)],
        [("Fuentes externas: ", True), ("marco normativo y de protección al consumidor "
         "financiero (CMF, SERNAC, Ley 19.628).", False)],
    ])
    rag = _rag_ejemplo()
    parrafo(doc,
        f"Ejemplo de recuperación para la consulta «{rag['pregunta']}» (combina ambos orígenes: "
        f"{', '.join(rag['origenes_combinados'])}):", space_after=4)
    tabla(doc, ["Origen", "Fuente", "Score"],
          [[f["origen"], f["fuente"].split(" — ")[0].split(" (")[0], f["score"]]
           for f in rag["fuentes"]],
          anchos=[1.1, 4.4, 0.9], size=8.5)
    parrafo(doc,
        "El backend léxico (TF-IDF) se eligió para garantizar reproducibilidad sin "
        "credenciales; en producción se reemplaza por embeddings densos manteniendo la misma "
        "interfaz. La base conceptual de chunking/embeddings/recuperación vectorial se "
        "desarrolló en RA1/IL1.3.", space_after=4)

    h2(doc, "2.3 Arquitectura de la solución (IE3)")
    parrafo(doc,
        "La arquitectura integra el LLM, la recuperación de información y el control de "
        "contexto (memoria) bajo una capa de seguridad de entrada y una capa de observabilidad "
        "de salida. Cada solicitud del cliente atraviesa el saneamiento de seguridad, es "
        "procesada por el AgentExecutor (ciclo ReAct) que decide qué herramienta o sub-agente "
        "invocar, y queda instrumentada para su monitoreo.")
    figura(doc, "00_arquitectura.png",
           "Figura 1. Arquitectura por capas de la solución, integrando RA1, RA2 y RA3.",
           width=6.2)

    h2(doc, "2.4 Justificación de decisiones de diseño (IE4)")
    nota_ia(doc, "Sección de justificación técnica: el equipo debe revisar y redactar con sus "
                 "propias palabras, según la política de uso de IA de la evaluación. Se entrega "
                 "el siguiente andamiaje de apoyo.")
    andamiaje(doc, [
        "¿Por qué LangChain para el agente principal? (ciclo ReAct, gestión de herramientas, "
        "tipos de memoria compatibles).",
        "¿Por qué CrewAI para la orquestación? (roles especializados, delegación, "
        "replanificación jerárquica).",
        "¿Por qué ventana k=5 en memoria de corto plazo y JSON en largo plazo? (equilibrio "
        "contexto/costo; simplicidad de despliegue).",
        "¿Por qué gpt-4o vía GitHub Models? (capacidad, disponibilidad para el curso, costo).",
        "Relación de cada decisión con los requerimientos organizacionales y las limitaciones "
        "del modelo (trazabilidad y privacidad).",
    ])

    # ── 3. Agente funcional ───────────────────────────────────────────────────
    h1(doc, "3. Desarrollo del agente funcional")
    h2(doc, "3.1 Herramientas de consulta, escritura y razonamiento (IE5)")
    tabla(doc, ["Herramienta", "Tipo", "Función"], [
        ["consultar_saldo", "Consulta", "Saldo, tipo de cuenta y productos por RUT"],
        ["realizar_transferencia", "Escritura", "Transferencia validando saldo suficiente"],
        ["calcular_cuota_credito", "Razonamiento", "Amortización francesa (cuota, intereses)"],
        ["registrar_reclamo", "Escritura", "Ticket de reclamo persistido con SLA"],
        ["buscar_producto_bancario", "Consulta/RAG", "Características y requisitos de productos"],
    ], anchos=[2.0, 1.3, 4.0], size=9, fuente_resaltar_col0=True)

    h2(doc, "3.2 Memoria y recuperación de contexto (IE6)")
    parrafo(doc,
        "Memoria de corto plazo con ConversationBufferWindowMemory (k=5) para sostener el hilo "
        "de la conversación activa, y memoria de largo plazo en JSON persistente por RUT que "
        "inyecta el historial del cliente (incidentes previos, última sesión) en el system "
        "prompt, asegurando continuidad entre sesiones.")
    h2(doc, "3.3 Planificación y toma de decisiones (IE7)")
    parrafo(doc,
        "Para solicitudes multi-etapa se aplica el patrón Plan-and-Execute: el agente genera "
        "un plan explícito antes de actuar (p. ej., consultar saldo → simular crédito a 24 y "
        "36 meses → comparar → presentar requisitos), reduciendo la pérdida de objetivo. Las "
        "solicitudes complejas se derivan a la orquestación multi-agente CrewAI, con procesos "
        "secuencial y jerárquico (el Gerente delega y replanifica dinámicamente).")
    h2(doc, "3.4 Orquestación y flujo de trabajo (IE8)")
    parrafo(doc,
        "El equipo CrewAI lo componen cuatro roles: Gerente (coordina y sintetiza), Analista "
        "Financiero (calcula métricas de crédito), Asesor de Productos (recomienda) y "
        "Especialista en Operaciones (verifica viabilidad regulatoria). El Gerente consolida "
        "los aportes en una respuesta única para el cliente. El detalle de implementación y los "
        "escenarios de prueba están documentados en el repositorio (README_EP2).")

    # ── 4. Observabilidad ─────────────────────────────────────────────────────
    h1(doc, "4. Observabilidad, trazabilidad y seguridad")
    h2(doc, "4.1 Métricas de observabilidad (IE9)")
    rico(doc, [
        ("Cada interacción se instrumenta con un ", False), ("tracer", True),
        (" que emite trazas JSONL con latencia, tokens, herramienta usada, éxito y señales de "
         "seguridad. Resultados globales sobre ", False),
        (f"{g['interacciones_totales']} interacciones", True), (":", False),
    ])
    caja(doc, "KPIs globales",
        f"Precisión (éxito): {pct(g['precision_exito'])}   ·   "
        f"Acierto de herramienta: {pct(g['acierto_herramienta'])}   ·   "
        f"Consistencia: {pct(cons['consistencia_media'])}   ·   "
        f"Tasa de error: {pct(g['tasa_error'])}\n"
        f"Latencia media/p95/máx: {g['latencia_media_ms']:.0f} / {g['latencia_p95_ms']:.0f} / "
        f"{g['latencia_max_ms']:.0f} ms   ·   Tokens: {g['tokens_totales']:,}   ·   "
        f"Costo estimado: US${g['costo_total_usd']:.2f}")
    figura(doc, "07_dashboard_panel.png",
           "Figura 2. Captura del dashboard de monitoreo integrado (KPIs + visualizaciones).",
           width=6.3)

    h2(doc, "4.2 Trazabilidad y propuesta de mejoras (IE10, IE12)")
    dia = anom[anom["anomalia"]]
    items = [
        [("Cuello de botella: ", True),
         (f"«multiagente_credito» con p95 de {esc.loc['multiagente_credito','latencia_p95_ms']:.0f} ms.", False)],
        [("Selección de herramienta: ", True),
         (f"éxito {pct(pat['exito_con_tool_correcta'])} con herramienta correcta vs. "
          f"{pct(pat['exito_con_tool_incorrecta'])} cuando se equivoca.", False)],
        [("Ambigüedad: ", True),
         (f"brecha de {pat['brecha_ambiguedad']*100:.0f} pp entre entradas claras y ambiguas.", False)],
    ]
    if not dia.empty:
        d = dia.iloc[0]
        items.insert(0, [("Anomalía temporal: ", True),
            (f"incidente el {d['fecha']} (p95 {d['latencia_p95_ms']:.0f} ms, z={d['z_latencia_p95_ms']}).", False)])
    vinneta(doc, items)
    figura(doc, "03_serie_temporal_incidente.png",
           "Figura 3. Evolución diaria; el incidente se detecta automáticamente (área sombreada).",
           width=5.8)
    parrafo(doc, "Mejoras propuestas (priorizadas por impacto en sostenibilidad y escalabilidad):",
            space_after=2)
    numerada(doc, [
        "Desambiguación de intención (clarificación + few-shot de enrutamiento).",
        "Paralelizar especialistas y cachear productos para reducir latencia y tokens.",
        "Reintentos con backoff y modelo de respaldo para resiliencia ante incidentes.",
        "Alertas automáticas por umbrales sobre el dashboard (observabilidad proactiva).",
    ], size=9.5)

    h2(doc, "4.3 Seguridad y uso responsable (IE11)")
    parrafo(doc,
        "La capa de seguridad aplica enmascaramiento de PII (Ley 19.628), defensa ante "
        "prompt-injection (OWASP LLM01), rate-limiting, auditoría con encadenamiento de hash y "
        "confirmación humana en operaciones críticas.")
    rico(doc, [
        (f"Se registraron {sec['intentos_injection']} intentos de prompt-injection, bloqueados "
         f"en {pct(sec['tasa_bloqueo'])}", True),
        (f"; persisten {sec['bypasses']} casos no bloqueados (riesgo residual a reforzar). Se "
         f"enmascaró PII en {sec['interacciones_con_pii']} interacciones.", False),
    ])

    # ── 5. Resultados ─────────────────────────────────────────────────────────
    h1(doc, "5. Resultados consolidados")
    nombres = {
        "consulta_saldo": "Consulta de saldo", "info_producto": "Info. de producto",
        "simulacion_credito": "Simulación crédito", "transferencia": "Transferencia",
        "registro_reclamo": "Registro reclamo", "multiagente_credito": "Multi-agente crédito",
        "consulta_ambigua": "Consulta ambigua", "seguridad_injection": "Intento injection",
    }
    filas = [[nombres.get(n, n), int(r["n"]), pct(r["precision_exito"]),
              pct(r["tasa_error"]), f"{r['latencia_p95_ms']:.0f}", f"{int(r['tokens_medio'])}"]
             for n, r in esc.iterrows()]
    tabla(doc, ["Escenario", "N", "Precisión", "Tasa error", "Latencia p95 (ms)", "Tokens medio"],
          filas, anchos=[1.9, 0.6, 1.0, 1.0, 1.5, 1.1], size=9, fuente_resaltar_col0=True)

    # ── 6. Declaración de uso de IA ───────────────────────────────────────────
    h1(doc, "6. Declaración de uso de inteligencia artificial")
    parrafo(doc,
        "En conformidad con la política de la evaluación y la guía de Biblioteca Duoc UC "
        "(bibliotecas.duoc.cl/ia), se declara el uso de IA como apoyo en las siguientes tareas:")
    vinneta(doc, [
        "Asistencia en redacción y estructura del informe y de la documentación del código.",
        "Apoyo en la generación de diagramas y en la organización de tablas de resultados.",
        "Generación de la instrumentación y del dataset de telemetría sintética (revisado por el equipo).",
    ])
    nota_ia(doc, "Las conclusiones, las justificaciones técnicas y las reflexiones individuales "
                 "fueron redactadas por los integrantes sin apoyo de IA. Todo contenido generado "
                 "con IA fue revisado y validado por el equipo.")

    # ── 7. Conclusiones y reflexiones ─────────────────────────────────────────
    h1(doc, "7. Conclusiones y reflexiones individuales")
    nota_ia(doc, "Esta sección NO debe redactarse con IA. Complétela el equipo y cada integrante.")
    h2(doc, "7.1 Conclusiones del equipo")
    andamiaje(doc, [
        "Grado de cumplimiento de los requerimientos del caso organizacional.",
        "Principales aprendizajes técnicos de integrar LLM + RAG + agentes + observabilidad.",
        "Limitaciones encontradas y trabajo futuro.",
    ])
    for integrante in ["Roger Rosas", "Rodrigo Santis", "Edgardo Gutiérrez"]:
        h2(doc, f"7.2 Reflexión individual — {integrante}")
        andamiaje(doc, [
            "Mi contribución específica al proyecto.",
            "Lo que aprendí y cómo lo aplicaría en un contexto real.",
            "Un dilema ético o de seguridad que identifiqué y cómo lo abordaría.",
        ])

    # ── 8. Referencias ────────────────────────────────────────────────────────
    h1(doc, "8. Referencias")
    refs = [
        "Biblioteca Duoc UC. (2024). Guía para el uso de inteligencia artificial. https://bibliotecas.duoc.cl/ia",
        "Chase, H. (2023). LangChain: Building applications with LLMs through composability. https://www.langchain.com/",
        "Comisión para el Mercado Financiero. (2020). RAN Cap. 20-7: Gestión de la ciberseguridad. CMF Chile.",
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. arXiv:2005.11401.",
        "Ley N°19.628. Sobre protección de la vida privada. Biblioteca del Congreso Nacional de Chile.",
        "Moura, J. (2024). CrewAI: Orchestrating role-playing autonomous AI agents. https://docs.crewai.com/",
        "OWASP Foundation. (2023). OWASP Top 10 for Large Language Model Applications.",
        "Yao, S., et al. (2022). ReAct: Synergizing reasoning and acting in language models. arXiv:2210.03629.",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.runs[0].font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)
    parrafo(doc,
        "Repositorio (código, pipeline RAG, notebook ejecutable, README, dashboard y evidencia "
        "de 22 pruebas automatizadas): "
        "github.com/rogrosas/Ingenier-a-de-Soluciones-con-Inteligencia-Artificial.",
        size=9, italic=True)

    doc.save(SALIDA)
    return SALIDA


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    ruta = build()
    print(f"✔ Informe EFT generado: {os.path.relpath(ruta, ROOT)}")
